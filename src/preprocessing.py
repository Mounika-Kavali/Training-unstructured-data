#preprocessing.py

from PyPDF2 import PdfReader
from PIL import Image
import os
import docx

def extract_text_from_pdf(pdf):
    text = ""
    pdf_reader = PdfReader(pdf)
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

def extract_text_from_docx(docx_file_path):
    doc = docx.Document(docx_file_path)
    
    # Extract text
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    
    # Extract tables
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        tables.append(table_data)
    
    return text, tables


def extract_images_from_pdf(pdf, image_folder):
    images = []
    pdf_reader = PdfReader(pdf)

    for page_number, page in enumerate(pdf_reader.pages):
        if '/XObject' in page:
            for img_number, img_ref in enumerate(page['/XObject']):
                img = page['/XObject'][img_ref]
                if img['/Subtype'] == '/Image':
                    img_data = img._data
                    img_obj = Image.frombytes(img['/ColorSpace'], img['/BitsPerComponent'], img_data)
                    img_path = os.path.join(image_folder, f"page{page_number + 1}_img{img_number + 1}.png")
                    img_obj.save(img_path)
                    images.append(img_path)

    return images